import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def set_cell_background(cell, fill_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_color}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'''
        <w:tcMar {nsdecls("w")}>
            <w:top w:w="{top}" w:type="dxa"/>
            <w:bottom w:w="{bottom}" w:type="dxa"/>
            <w:left w:w="{left}" w:type="dxa"/>
            <w:right w:w="{right}" w:type="dxa"/>
        </w:tcMar>
    ''')
    tcPr.append(tcMar)

def add_code_block(doc, code_text):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    
    cell = tbl.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_background(cell, "F4F6F8")
    set_cell_margins(cell, top=120, bottom=120, left=180, right=180)
    
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'''
        <w:tcBorders {nsdecls("w")}>
            <w:top w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>
            <w:left w:val="single" w:sz="24" w:space="0" w:color="0066CC"/>
            <w:bottom w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>
            <w:right w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>
        </w:tcBorders>
    ''')
    tcPr.append(tcBorders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x24, 0x29, 0x2E)
    
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

def create_document():
    doc = docx.Document()
    
    # Page setup - Margins 1 inch
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_title.paragraph_format.space_after = Pt(2)
    run_title = p_title.add_run("Manual Completo de Conexão ao Banco de Dados")
    run_title.font.name = 'Segoe UI'
    run_title.font.size = Pt(24)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(0x00, 0x40, 0x80)
    
    # Subtitle
    p_sub = doc.add_paragraph()
    p_sub.paragraph_format.space_after = Pt(18)
    run_sub = p_sub.add_run("Projeto Portal Turismo SP (TravelGPT-V1) | PostgreSQL & Vercel Serverless")
    run_sub.font.name = 'Segoe UI'
    run_sub.font.size = Pt(13)
    run_sub.font.italic = True
    run_sub.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    
    # Section 1
    h1 = doc.add_heading("1. Visão Geral da Arquitetura de Banco de Dados", level=1)
    h1.style.font.name = 'Segoe UI'
    h1.style.font.color.rgb = RGBColor(0x00, 0x40, 0x80)
    
    p = doc.add_paragraph(
        "Este documento contém o guia definitivo e passo a passo para configurar, conectar, "
        "migrar e consultar o banco de dados PostgreSQL no projeto Portal Turismo SP (TravelGPT-V1). "
        "A aplicação é construída sobre uma arquitetura moderna baseada em Node.js, Vite (Frontend), "
        "Serverless Functions da Vercel (Backend API) e banco de dados relacional PostgreSQL (Vercel Postgres)."
    )
    p.paragraph_format.space_after = Pt(10)
    
    # Summary Table
    table = doc.add_table(rows=5, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    headers = ["Componente", "Tecnologia / Especificação"]
    for j, text in enumerate(headers):
        cell = table.cell(0, j)
        set_cell_background(cell, "004080")
        set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        
    data = [
        ("Motor de Banco de Dados", "PostgreSQL 15+ (Vercel Postgres / Neon / Supabase)"),
        ("Driver Node.js", "pg (node-postgres) v8.13.1 & @vercel/postgres v0.10.0"),
        ("Script de Migração / Seed", "node scripts/seed_database.js (npm run db:seed)"),
        ("API Serverless Endpoint", "api/places.js (/api/places)")
    ]
    
    for i, row in enumerate(data):
        for j, text in enumerate(row):
            cell = table.cell(i+1, j)
            bg = "F9FAFB" if i % 2 == 0 else "FFFFFF"
            set_cell_background(cell, bg)
            set_cell_margins(cell, top=80, bottom=80, left=150, right=150)
            p = cell.paragraphs[0]
            p.add_run(text)
            
    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    
    # Section 2
    h2 = doc.add_heading("2. Pré-requisitos", level=1)
    h2.style.font.name = 'Segoe UI'
    h2.style.font.color.rgb = RGBColor(0x00, 0x40, 0x80)
    
    p = doc.add_paragraph("Antes de iniciar a configuração da conexão, garanta que seu ambiente cumpra os seguintes requisitos:")
    p.paragraph_format.space_after = Pt(6)
    
    reqs = [
        "Node.js instalado (versão 18.x ou superior).",
        "Gerenciador de pacotes NPM (incluso no Node.js).",
        "Instância ativa de um banco PostgreSQL (na Vercel Postgres, Supabase, Neon ou local).",
        "CLI da Vercel instalada opcionalmente (npx vercel) para sincronização automática de variáveis."
    ]
    for r in reqs:
        bp = doc.add_paragraph(style='List Bullet')
        bp.paragraph_format.space_after = Pt(4)
        bp.add_run(r)
        
    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    
    # Section 3
    h3 = doc.add_heading("3. Passo a Passo da Configuração das Variáveis de Ambiente", level=1)
    h3.style.font.name = 'Segoe UI'
    h3.style.font.color.rgb = RGBColor(0x00, 0x40, 0x80)
    
    p = doc.add_paragraph(
        "A aplicação lida de forma resiliente com as variáveis de conexão. Tanto o script de seed "
        "quanto o endpoint de API leem o banco através das variáveis de ambiente na seguinte ordem de prioridade:"
    )
    p.paragraph_format.space_after = Pt(6)
    
    add_code_block(doc, "const connectionString = process.env.POSTGRES_URL || process.env.DATABASE_URL || process.env.POSTGRES_URL_NON_POOLING;")
    
    h3_1 = doc.add_heading("3.1 Método 1: Sincronização Automática via Vercel CLI (Recomendado)", level=2)
    h3_1.style.font.name = 'Segoe UI'
    h3_1.style.font.color.rgb = RGBColor(0x00, 0x66, 0xCC)
    p = doc.add_paragraph(
        "Se o seu repositório já está vinculado a um projeto na Vercel com a extensão Vercel Postgres ativada, "
        "você pode puxar automaticamente todas as credenciais de produção/desenvolvimento executando:"
    )
    p.paragraph_format.space_after = Pt(4)
    add_code_block(doc, "# 1. Faça login na Vercel (se ainda não estiver autenticado)\nnpx vercel login\n\n# 2. Vincule o projeto local ao projeto na Vercel\nnpx vercel link\n\n# 3. Baixe as variáveis de ambiente diretamente para o arquivo .env.local\nnpx vercel env pull .env.local")
    
    h3_2 = doc.add_heading("3.2 Método 2: Configuração Manual do Arquivo .env.local", level=2)
    h3_2.style.font.name = 'Segoe UI'
    h3_2.style.font.color.rgb = RGBColor(0x00, 0x66, 0xCC)
    p = doc.add_paragraph(
        "Caso esteja utilizando uma instância local do PostgreSQL ou outro provedor de nuvem (como Neon ou Supabase), "
        "crie o arquivo .env.local na raiz do projeto (d:\\PROJETOS\\ANTIGRAVITY\\TRAVELGPT-V1\\.env.local) e adicione:"
    )
    p.paragraph_format.space_after = Pt(4)
    add_code_block(doc, '# Arquivo: .env.local\nPOSTGRES_URL="postgres://usuario:senha@host.postgres.vercel-storage.com:5432/verceldb?sslmode=require"')
    
    p = doc.add_paragraph("Estrutura detalhada da String de Conexão (URL):")
    p.paragraph_format.space_after = Pt(4)
    
    url_parts = [
        ("postgres://", "Protocolo de comunicação do PostgreSQL."),
        ("usuario", "Nome do usuário cadastrado no banco de dados."),
        ("senha", "Senha de acesso do usuário."),
        ("host", "Endereço IP ou hostname do servidor de banco de dados."),
        ("5432", "Porta padrão de comunicação do PostgreSQL."),
        ("verceldb", "Nome da base de dados selecionada."),
        ("?sslmode=require", "Parâmetro obrigatório para conexões seguras criptografadas (SSL/TLS).")
    ]
    
    for term, desc in url_parts:
        bp = doc.add_paragraph(style='List Bullet')
        bp.paragraph_format.space_after = Pt(3)
        r1 = bp.add_run(f"{term}: ")
        r1.bold = True
        bp.add_run(desc)
        
    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    
    # Section 4
    h4 = doc.add_heading("4. Modelagem de Dados e Schema SQL", level=1)
    h4.style.font.name = 'Segoe UI'
    h4.style.font.color.rgb = RGBColor(0x00, 0x40, 0x80)
    
    p = doc.add_paragraph(
        "O banco de dados é estruturado de forma relacional e otimizada para consultas de pontos turísticos, "
        "cidades e categorias. Abaixo está a definição SQL exata executada durante o provisionamento do banco:"
    )
    p.paragraph_format.space_after = Pt(6)
    
    add_code_block(doc, """-- 1. Tabela de Cidades
CREATE TABLE IF NOT EXISTS cities (
  id BIGSERIAL PRIMARY KEY,
  name VARCHAR(100) NOT NULL UNIQUE,
  state CHAR(2) NOT NULL DEFAULT 'SP',
  created_at TIMESTAMP WITH TIME ZONE DEFAULT CURRENT_TIMESTAMP
);

-- 2. Tabela de Categorias
CREATE TABLE IF NOT EXISTS categories (
  id BIGSERIAL PRIMARY KEY,
  name VARCHAR(100) NOT NULL UNIQUE,
  created_at TIMESTAMP WITH TIME ZONE DEFAULT CURRENT_TIMESTAMP
);

-- 3. Tabela de Locais e Pontos Turísticos
CREATE TABLE IF NOT EXISTS places (
  id VARCHAR(50) PRIMARY KEY,
  title VARCHAR(255) NOT NULL,
  category_id BIGINT REFERENCES categories(id) ON DELETE SET NULL,
  original_category VARCHAR(100),
  city_id BIGINT REFERENCES cities(id) ON DELETE RESTRICT,
  address VARCHAR(500) NOT NULL,
  phone VARCHAR(50),
  email VARCHAR(255),
  google_maps_url TEXT,
  latitude DECIMAL(10, 8) NOT NULL,
  longitude DECIMAL(11, 8) NOT NULL,
  tier VARCHAR(20) DEFAULT 'bronze',
  is_featured BOOLEAN NOT NULL DEFAULT FALSE,
  cover_image_url TEXT,
  backup_image_url TEXT,
  description TEXT,
  rating DECIMAL(3, 1) DEFAULT 4.2,
  created_at TIMESTAMP WITH TIME ZONE DEFAULT CURRENT_TIMESTAMP
);""")

    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    
    # Section 5
    h5 = doc.add_heading("5. Execução de Migração e Povoamento (Seed)", level=1)
    h5.style.font.name = 'Segoe UI'
    h5.style.font.color.rgb = RGBColor(0x00, 0x40, 0x80)
    
    p = doc.add_paragraph(
        "O projeto conta com um script automatizado (scripts/seed_database.js) que lê o arquivo "
        "src/data/places.json, cria as tabelas se não existirem, resolve dinamicamente as chaves estrangeiras "
        "de cidades e categorias e faz o upsert de todos os pontos turísticos."
    )
    p.paragraph_format.space_after = Pt(6)
    
    p = doc.add_paragraph("Para executar a migração e popular o banco, rode o seguinte comando no terminal:")
    p.paragraph_format.space_after = Pt(4)
    add_code_block(doc, "npm run db:seed")
    
    p = doc.add_paragraph("Saída esperada em caso de sucesso:")
    p.paragraph_format.space_after = Pt(4)
    add_code_block(doc, """🚀 Iniciando criação das tabelas e migração de dados no Vercel Postgres...
📦 Carregados 45 registros de places.json
✅ Sucesso! Migração concluída com 45 pontos turísticos inseridos/atualizados na base Postgres no Vercel.""")

    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    
    # Section 6
    h6 = doc.add_heading("6. Arquitetura da API Serverless (api/places.js)", level=1)
    h6.style.font.name = 'Segoe UI'
    h6.style.font.color.rgb = RGBColor(0x00, 0x40, 0x80)
    
    p = doc.add_paragraph(
        "A consulta dos dados no frontend é realizada através da Serverless Function api/places.js. "
        "Ela estabelece conexão via Pool da biblioteca 'pg', lida com requisições CORS e executa "
        "uma query otimizada com JOINs para retornar o formato exato esperado pela interface React."
    )
    p.paragraph_format.space_after = Pt(6)
    
    add_code_block(doc, """import pg from 'pg';

const connectionString = process.env.POSTGRES_URL || process.env.DATABASE_URL || process.env.POSTGRES_URL_NON_POOLING;

let pool;
if (connectionString) {
  const { Pool } = pg;
  pool = new Pool({
    connectionString,
    ssl: connectionString.includes('localhost') ? false : { rejectUnauthorized: false }
  });
}

export default async function handler(req, res) {
  // CORS & Validação de Métodos HTTP...
  const client = await pool.connect();
  try {
    const query = `
      SELECT 
        p.id, p.title, c.name AS category, p.original_category AS "originalCategory",
        ci.name AS city, p.address, p.phone, p.email, p.google_maps_url AS "googleMapsUrl",
        p.latitude::float AS lat, p.longitude::float AS lng, p.tier,
        p.is_featured AS "isFeatured", p.cover_image_url AS "coverImage",
        p.backup_image_url AS "backupImage", p.description, p.rating::float AS rating
      FROM places p
      LEFT JOIN categories c ON p.category_id = c.id
      LEFT JOIN cities ci ON p.city_id = ci.id
      ORDER BY p.is_featured DESC, p.title ASC;
    `;
    const result = await client.query(query);
    return res.status(200).json(result.rows);
  } finally {
    client.release();
  }
}""")

    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    
    # Section 7
    h7 = doc.add_heading("7. Diagnóstico e Resolução de Problemas (Troubleshooting)", level=1)
    h7.style.font.name = 'Segoe UI'
    h7.style.font.color.rgb = RGBColor(0x00, 0x40, 0x80)
    
    issues = [
        ("❌ ERRO: Nenhuma variável de conexão encontrada", 
         "Causa: O arquivo .env.local não existe ou a variável POSTGRES_URL não está definida.\n"
         "Solução: Execute 'npx vercel env pull .env.local' ou crie o arquivo .env.local manualmente com a string POSTGRES_URL."),
         
        ("🔒 Erro de SSL / Self-signed certificate", 
         "Causa: Conexão segura exigida pelo provedor em nuvem sem suporte a certificados locais estritos.\n"
         "Solução: O código já trata isso configurando ssl: { rejectUnauthorized: false } para servidores remotos e desativando SSL apenas para localhost."),
         
        ("⚡ Erro de Limite de Conexões (Connection Pool Exhausted)", 
         "Causa: Serverless functions abrindo muitas conexões simultâneas sem fechar.\n"
         "Solução: O código utiliza o bloco 'finally { client.release(); }' garantindo a liberação imediata da conexão ao Pool após a consulta.")
    ]
    
    for title, detail in issues:
        p_t = doc.add_paragraph()
        p_t.paragraph_format.space_after = Pt(2)
        r = p_t.add_run(title)
        r.bold = True
        r.font.color.rgb = RGBColor(0xA0, 0x00, 0x00)
        
        p_d = doc.add_paragraph(detail)
        p_d.paragraph_format.space_after = Pt(8)
        p_d.paragraph_format.left_indent = Inches(0.2)
        
    doc.add_paragraph().paragraph_format.space_after = Pt(18)
    
    # Save document
    output_path = r"d:\PROJETOS\ANTIGRAVITY\TRAVELGPT-V1\Instrucoes_Conexao_Banco_de_Dados.docx"
    doc.save(output_path)
    print(f"Documento criado com sucesso em: {output_path}")

if __name__ == "__main__":
    create_document()
